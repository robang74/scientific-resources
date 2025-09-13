#!/usr/bin/env python3
"""
DOCX Review Splitter - Clean CLI Version
A command-line tool to split DOCX files containing multiple reviews into individual review files.

Usage:
    python docx_splitter.py <input_file> [output_dir]
    python docx_splitter.py --all [output_dir]
    python docx_splitter.py --help

Examples:
    python docx_splitter.py "Reviews_1-30.docx" split-reviews/
    python docx_splitter.py --all
"""

import os
import sys
import argparse
from docx import Document
import re
from collections import defaultdict
from pathlib import Path

class DocxSplitter:
    """Main class for splitting DOCX files into individual reviews."""
    
    def __init__(self, verbose=True):
        self.verbose = verbose
        
    def log(self, message):
        """Print message if verbose mode is enabled."""
        if self.verbose:
            print(message)
    
    def find_review_boundaries(self, doc):
        """
        Find all review boundaries in the document.
        Handles various formats: "Review X:", "Review X, Short:", etc.
        
        Returns:
            list: List of tuples (start_idx, end_idx, review_number, title, suffix)
        """
        boundaries = []
        review_counts = defaultdict(int)
        
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            
            # Look for any line starting with "Review" followed by a number
            match = re.search(r'Review\s+(\d+)', text, re.IGNORECASE)
            
            if match and text.startswith('Review'):
                review_number = int(match.group(1))
                
                # Extract title from the rest of the line
                title = text[match.end():].strip()
                
                # Clean up title (remove colons, commas, handle "Short")
                title = re.sub(r'^[,:]\s*', '', title)
                title = re.sub(r'^Short[,:]\s*', '[Short] ', title, flags=re.IGNORECASE)
                
                # If title is empty or very short, look at next paragraph
                if not title or len(title) < 5:
                    if i + 1 < len(doc.paragraphs):
                        next_text = doc.paragraphs[i + 1].text.strip()
                        if next_text and not next_text.startswith('http'):
                            title = next_text
                
                # Handle duplicate review numbers (e.g., Review 117 appears twice)
                review_counts[review_number] += 1
                suffix = ""
                if review_counts[review_number] > 1:
                    suffix = chr(ord('a') + review_counts[review_number] - 1)  # a, b, c, etc.
                
                boundaries.append((i, None, review_number, title, suffix))
        
        # Set end boundaries (each review ends where the next begins)
        for i in range(len(boundaries)):
            if i < len(boundaries) - 1:
                end_idx = boundaries[i+1][0] - 1
                boundaries[i] = (boundaries[i][0], end_idx, boundaries[i][2], boundaries[i][3], boundaries[i][4])
            else:
                # Last review goes to end of document
                end_idx = len(doc.paragraphs) - 1
                boundaries[i] = (boundaries[i][0], end_idx, boundaries[i][2], boundaries[i][3], boundaries[i][4])
        
        return boundaries
    
    def extract_review(self, doc, start_idx, end_idx, review_number, title, suffix):
        """
        Extract a single review and create a new DOCX document.
        
        Returns:
            Document: New DOCX document containing the review, or None if extraction fails
        """
        try:
            new_doc = Document()
            
            # Add title with suffix if needed
            display_number = f"{review_number}{suffix}" if suffix else str(review_number)
            new_doc.add_heading(f"Review {display_number}: {title}", level=1)
            
            # Copy paragraphs from original document
            for i in range(start_idx, min(end_idx + 1, len(doc.paragraphs))):
                if i == start_idx:
                    continue  # Skip the original title paragraph
                
                original_para = doc.paragraphs[i]
                new_para = new_doc.add_paragraph()
                new_para.text = original_para.text
                
                # Preserve formatting if possible
                try:
                    new_para.style = original_para.style
                except:
                    pass  # Use default if style doesn't exist
            
            return new_doc
            
        except Exception as e:
            if self.verbose:
                print(f"Error extracting review {review_number}{suffix}: {str(e)}")
            return None
    
    def split_file(self, input_path, output_dir="split-reviews"):
        """
        Split a single DOCX file into individual review files.
        
        Args:
            input_path (str): Path to input DOCX file
            output_dir (str): Output directory for individual review files
            
        Returns:
            int: Number of reviews successfully extracted
        """
        input_path = Path(input_path)
        output_dir = Path(output_dir)
        
        if not input_path.exists():
            self.log(f"Error: Input file '{input_path}' not found!")
            return 0
        
        self.log(f"\n=== Splitting: {input_path.name} ===")
        
        try:
            # Load the document
            doc = Document(input_path)
            
            # Find review boundaries
            boundaries = self.find_review_boundaries(doc)
            self.log(f"Found {len(boundaries)} reviews to extract")
            
            if not boundaries:
                self.log("No reviews found!")
                return 0
            
            # Create output directory
            output_dir.mkdir(parents=True, exist_ok=True)
            
            # Extract each review
            success_count = 0
            for start_idx, end_idx, review_num, title, suffix in boundaries:
                display_number = f"{review_num}{suffix}" if suffix else str(review_num)
                
                if self.verbose:
                    print(f"Extracting Review {display_number}: {title[:50]}...")
                
                # Extract the review
                review_doc = self.extract_review(doc, start_idx, end_idx, review_num, title, suffix)
                
                if review_doc:
                    # Save the review
                    filename = f"Review_{review_num:03d}{suffix}.docx"
                    output_path = output_dir / filename
                    review_doc.save(str(output_path))
                    
                    if self.verbose:
                        print(f"  ✅ Saved: {filename}")
                    success_count += 1
                else:
                    if self.verbose:
                        print(f"  ❌ Failed to extract Review {review_num}{suffix}")
            
            self.log(f"\n=== Splitting Complete ===")
            self.log(f"Successfully extracted: {success_count}/{len(boundaries)} reviews")
            return success_count
            
        except Exception as e:
            self.log(f"Error processing {input_path}: {str(e)}")
            return 0
    
    def split_all_files(self, docx_dir="mike-paper-reviews-500/docx", output_dir="split-reviews"):
        """
        Split all review DOCX files in a directory.
        
        Args:
            docx_dir (str): Directory containing DOCX files
            output_dir (str): Output directory for individual review files
            
        Returns:
            int: Total number of reviews extracted
        """
        docx_dir = Path(docx_dir)
        
        if not docx_dir.exists():
            self.log(f"Error: DOCX directory '{docx_dir}' not found!")
            return 0
        
        # Define the review files in order
        review_files = [
            "Reviews_1-30.docx",
            "Reviews_31-60.docx", 
            "Reviews_61-90.docx",
            "Reviews_91-120.docx",
            "Reviews_121-150.docx",
            "Reviews_151-180.docx",
            "Reviews_181-207.docx"
        ]
        
        self.log("DOCX Review Splitter - Processing All Files")
        self.log("=" * 50)
        
        total_extracted = 0
        
        for filename in review_files:
            file_path = docx_dir / filename
            if file_path.exists():
                extracted = self.split_file(file_path, output_dir)
                total_extracted += extracted
            else:
                self.log(f"Warning: {filename} not found, skipping...")
        
        self.log(f"\n🎉 ALL FILES PROCESSED!")
        self.log(f"�� Total reviews extracted: {total_extracted}")
        
        return total_extracted

def main():
    """Main CLI function."""
    parser = argparse.ArgumentParser(
        description="Split DOCX files containing multiple reviews into individual review files",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python docx_splitter.py Reviews_1-30.docx
  python docx_splitter.py Reviews_1-30.docx my-output/
  python docx_splitter.py --all
  python docx_splitter.py --all custom-output/
  python docx_splitter.py --quiet Reviews_1-30.docx
        """
    )
    
    parser.add_argument('input_file', nargs='?', 
                       help='Input DOCX file to split')
    parser.add_argument('output_dir', nargs='?', default='split-reviews',
                       help='Output directory (default: split-reviews)')
    parser.add_argument('--all', action='store_true',
                       help='Process all review files in mike-paper-reviews-500/docx/')
    parser.add_argument('--docx-dir', default='mike-paper-reviews-500/docx',
                       help='Directory containing DOCX files (default: mike-paper-reviews-500/docx)')
    parser.add_argument('--quiet', '-q', action='store_true',
                       help='Run in quiet mode (minimal output)')
    
    args = parser.parse_args()
    
    # Create splitter instance
    splitter = DocxSplitter(verbose=not args.quiet)
    
    if args.all:
        # Process all files
        total = splitter.split_all_files(args.docx_dir, args.output_dir)
        if not args.quiet:
            print(f"\nProcessing complete. {total} reviews extracted.")
    elif args.input_file:
        # Process single file
        count = splitter.split_file(args.input_file, args.output_dir)
        if not args.quiet:
            print(f"\nProcessing complete. {count} reviews extracted.")
    else:
        # No arguments provided
        parser.print_help()
        sys.exit(1)

if __name__ == "__main__":
    main()
